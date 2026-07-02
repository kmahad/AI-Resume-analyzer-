import os
import pdfplumber
import docx
import io

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes using pdfplumber."""
    text_content = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_content.append(page_text)
    return "\n".join(text_content)

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes using python-docx, including tables."""
    text_content = []
    doc = docx.Document(io.BytesIO(file_bytes))
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_content.append(para.text)
            
    # Extract from tables (often used for layouts in resumes)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_content.append(" | ".join(row_text))
                
    return "\n".join(text_content)

def extract_text(file_bytes: bytes, filename: str) -> str:
    """Detect file type by extension and extract text, with a fallback for raw text."""
    ext = os.path.splitext(filename.lower())[1]
    
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        try:
            return extract_text_from_docx(file_bytes)
        except Exception as e:
            # If python-docx fails, let's attempt to read as zip and extract raw text from document.xml
            # or raise the exception. Raising/logging is fine since python-docx is installed.
            raise ValueError(f"Failed to parse Word document: {str(e)}")
    elif ext in [".txt", ".md"]:
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1")
    else:
        # Generic fallback
        try:
            return file_bytes.decode("utf-8")
        except Exception:
            raise ValueError(f"Unsupported file format: {ext}")
